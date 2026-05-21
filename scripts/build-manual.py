"""계측기 관리 사용 매뉴얼 (Word) 생성 스크립트.

실행:
    python scripts/build-manual.py

출력:
    docs/계측기-관리-매뉴얼.docx

이미지는 [이미지 N: 설명] 형식의 자리표시자로 들어감.
사용자가 Word에서 해당 위치에 캡처 이미지를 직접 삽입.
"""
from __future__ import annotations

import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


KOREAN_FONT = "맑은 고딕"


def set_korean_font(run, size_pt: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    """run에 한글/영문 폰트 일관 적용."""
    run.font.name = KOREAN_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rFonts.set(qn("w:ascii"), KOREAN_FONT)
    rFonts.set(qn("w:hAnsi"), KOREAN_FONT)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_korean_font(run, size_pt=24, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
    doc.add_paragraph()  # spacer


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_korean_font(run, size_pt=16, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_korean_font(run, size_pt=13, bold=True, color=RGBColor(0x2A, 0x4B, 0x82))


def add_paragraph(doc: Document, text: str, indent_cm: float = 0.0) -> None:
    p = doc.add_paragraph()
    if indent_cm > 0:
        p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_korean_font(run, size_pt=11)


def add_step(doc: Document, num: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    num_run = p.add_run(f"{num}. ")
    set_korean_font(num_run, size_pt=11, bold=True, color=RGBColor(0x2A, 0x4B, 0x82))
    text_run = p.add_run(text)
    set_korean_font(text_run, size_pt=11)


def add_bullet(doc: Document, text: str, indent_cm: float = 0.8) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(2)
    bullet_run = p.add_run("• ")
    set_korean_font(bullet_run, size_pt=11, color=RGBColor(0x2A, 0x4B, 0x82))
    text_run = p.add_run(text)
    set_korean_font(text_run, size_pt=11)


def add_image_placeholder(doc: Document, label: str) -> None:
    """이미지 자리표시자 — 사용자가 Word에서 직접 캡처 삽입.

    회색 배경의 박스 형태로 시각적 구분.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # 회색 음영
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EEEEEE")
    pPr.append(shd)

    # 박스 테두리
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "dashed")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "999999")
        pBdr.append(b)
    pPr.append(pBdr)

    run = p.add_run(f"[ {label} ]\n(여기에 캡처 이미지를 삽입하세요)")
    set_korean_font(run, size_pt=10, color=RGBColor(0x66, 0x66, 0x66))


def add_note(doc: Document, text: str) -> None:
    """노란 음영 노트 박스."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)

    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFF8E1")
    pPr.append(shd)

    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "F9A825")
    pBdr.append(left)
    pPr.append(pBdr)

    icon_run = p.add_run("💡 ")
    set_korean_font(icon_run, size_pt=11, bold=True)
    text_run = p.add_run(text)
    set_korean_font(text_run, size_pt=10)


def add_divider(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build() -> str:
    doc = Document()

    # 페이지 여백 살짝 줄이기
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # 기본 스타일 한글 폰트 적용
    style = doc.styles["Normal"]
    style.font.name = KOREAN_FONT
    style_rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rFonts.set(qn("w:ascii"), KOREAN_FONT)
    rFonts.set(qn("w:hAnsi"), KOREAN_FONT)
    style_rPr.append(rFonts)

    # ── 표지 ──
    add_title(doc, "계측기 관리 사용 매뉴얼")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("에델예약 PWA — 계측기 관리 기능")
    set_korean_font(run, size_pt=12, color=RGBColor(0x55, 0x55, 0x55))

    add_divider(doc)

    # ── 1. 개요 ──
    add_h1(doc, "1. 개요")
    add_paragraph(
        doc,
        "계측기 관리 페이지는 사내 계측기의 사용 등록과 현재 사용 현황 조회를 통합 관리하는 기능입니다. "
        "사용 등록 시 Supabase DB와 구글 시트 ‘계측기 관리대장’의 사용부서 컬럼이 자동으로 갱신됩니다.",
    )

    add_h2(doc, "1.1 화면 진입")
    add_step(doc, 1, "에델예약 앱 하단 탭에서 ‘계측기’ 아이콘을 누릅니다.")
    add_step(doc, 2, "‘계측기 관리’ 화면이 표시됩니다.")
    add_image_placeholder(doc, "이미지 1 · 계측기 관리 메인 화면")

    add_h2(doc, "1.2 화면 구성")
    add_bullet(doc, "사용할 계측기 선택 — 사용 시작할 계측기를 검색해 선택하고 ‘사용’ 버튼으로 기록")
    add_bullet(doc, "계측기 찾기 — 특정 계측기가 현재 누구에 의해 언제 사용 중인지 조회")
    add_bullet(doc, "계측기 관리 대장 — 구글 시트 원본을 새 탭에서 열기")

    add_divider(doc)

    # ── 2. 계측기 사용 등록 ──
    add_h1(doc, "2. 계측기 사용 등록하기")
    add_paragraph(
        doc,
        "사용하려는 계측기를 검색하고, 자동완성 목록에서 정확한 항목을 선택한 뒤 ‘사용’ 버튼을 누르면 "
        "DB와 시트에 동시에 기록됩니다.",
    )

    add_h2(doc, "2.1 검색어 입력")
    add_step(doc, 1, "‘사용할 계측기 선택’ 섹션의 입력칸에 계측기명(한글 또는 영문)이나 모델명 일부를 입력합니다.")
    add_step(doc, 2, "0.2초 후 자동완성 목록이 입력칸 아래에 나타납니다 (최대 10건).")
    add_image_placeholder(doc, "이미지 2 · 사용 자동완성 드롭다운")

    add_h2(doc, "2.2 정확한 항목 선택")
    add_paragraph(
        doc,
        "같은 이름의 계측기가 여러 개 있을 수 있으므로, 부가 정보(영문명 · 모델명 · 기기번호)를 확인해 "
        "정확한 항목을 클릭하세요.",
    )
    add_note(
        doc,
        "예: ‘오실로스코프’로 검색하면 5개 행이 나옵니다. 영문명·모델·기기번호 조합으로 "
        "사용하려는 계측기를 정확히 식별하세요. 직접 입력만 한 경우 ‘사용’ 버튼이 비활성 상태로 유지됩니다.",
    )

    add_h2(doc, "2.3 사용 등록 완료")
    add_step(doc, 1, "선택된 항목이 입력칸 아래에 ‘선택됨 · …’ 형태로 요약 표시됩니다.")
    add_step(doc, 2, "‘사용’ 버튼이 활성화됩니다. 클릭하면 DB에 사용 이력이 추가되고 시트의 ‘사용부서’ 컬럼이 갱신됩니다.")
    add_step(doc, 3, "성공 시 ‘XXX 사용 등록 완료’ 메시지가 표시되고, 입력칸은 초기화됩니다.")

    add_note(
        doc,
        "시트 저장에 실패할 경우 ‘사용 등록은 됐지만 시트 업데이트 실패: …’ 메시지가 표시됩니다. "
        "이 경우 DB에는 정상 기록된 상태이며, 시트 측만 별도로 확인이 필요합니다.",
    )

    add_divider(doc)

    # ── 3. 계측기 찾기 ──
    add_h1(doc, "3. 사용 중 계측기 찾기")
    add_paragraph(
        doc,
        "특정 계측기를 누가 언제 사용 중인지 확인하는 기능입니다. "
        "이름 · 모델명 · 관리번호 어느 것으로도 검색 가능합니다.",
    )

    add_h2(doc, "3.1 검색어 입력")
    add_step(doc, 1, "‘계측기 찾기’ 섹션의 입력칸에 검색어를 입력합니다.")
    add_step(doc, 2, "자동완성 목록이 표시됩니다 (최대 8건).")
    add_image_placeholder(doc, "이미지 3 · 찾기 자동완성 드롭다운")

    add_h2(doc, "3.2 항목 선택")
    add_paragraph(
        doc,
        "사용 등록 화면과 동일하게, 같은 이름이 여러 개일 때는 부가 정보(영문명·모델·기기번호)로 "
        "원하는 항목을 정확히 식별해 클릭하세요.",
    )
    add_note(
        doc,
        "선택 후 입력칸 아래에 ‘선택됨 · … · 관리번호’ 형태로 표시됩니다. 관리번호까지 명시되어 "
        "어느 계측기를 선택했는지 한눈에 확인할 수 있습니다.",
    )
    add_image_placeholder(doc, "이미지 4 · 찾기 항목 선택 후 상태")

    add_h2(doc, "3.3 결과 확인")
    add_step(doc, 1, "‘찾기’ 버튼을 클릭합니다.")
    add_step(doc, 2, "사용 이력이 있는 경우: ‘YYYY년 M월 D일 OOO님이 사용중입니다.’가 표시됩니다.")
    add_step(doc, 3, "사용 이력이 없는 경우: ‘— 사용 기록이 없습니다.’가 표시됩니다.")
    add_image_placeholder(doc, "이미지 5 · 사용중 결과 표시")

    add_divider(doc)

    # ── 4. 계측기 관리 대장 ──
    add_h1(doc, "4. 계측기 관리 대장 열기")
    add_paragraph(
        doc,
        "구글 시트 원본 ‘계측기 관리대장’을 새 탭에서 직접 열어 전체 데이터를 확인하거나 편집할 수 있습니다.",
    )
    add_step(doc, 1, "화면 하단의 ‘계측기 관리 대장’ 버튼을 누릅니다.")
    add_step(doc, 2, "구글 시트가 새 탭으로 열립니다.")
    add_note(
        doc,
        "시트의 ‘사용부서’ 컬럼은 사용 등록 시 자동 갱신되므로, 시트에서 수동으로 편집할 필요는 없습니다.",
    )

    add_divider(doc)

    # ── 5. 참고 ──
    add_h1(doc, "5. 참고 사항")

    add_h2(doc, "5.1 동일 이름 계측기 식별 방법")
    add_paragraph(
        doc,
        "‘오실로스코프’처럼 동일 이름을 가진 계측기가 여러 대 있을 때는 다음 순서로 식별하세요:",
    )
    add_bullet(doc, "1순위: 영문명 (예: DIGITAL STORAGE OSCILLOSCOPE)")
    add_bullet(doc, "2순위: 모델명 (예: Protek 5100)")
    add_bullet(doc, "3순위: 기기번호 — 각 계측기 본체에 부착된 시리얼 번호와 일치하는 값")

    add_h2(doc, "5.2 자동완성이 표시되지 않을 때")
    add_bullet(doc, "검색어를 더 정확하게 입력해 보세요 (한글/영문 모두 가능).")
    add_bullet(doc, "관리번호로 검색 가능한 곳은 ‘계측기 찾기’만 해당합니다. ‘사용할 계측기 선택’은 이름/모델만 가능.")
    add_bullet(doc, "여전히 안 나오면 시트(‘계측기 관리 대장’)에 해당 계측기가 등록되어 있는지 확인하세요.")

    add_h2(doc, "5.3 사용 등록 후 화면이 안 바뀔 때")
    add_paragraph(
        doc,
        "‘사용 등록 완료’ 메시지가 표시되면 DB에는 정상 반영된 상태입니다. "
        "다른 사용자가 같은 시점에 찾기를 해도 결과가 즉시 보입니다. "
        "본인 화면이 이상하면 화면을 위로 끌어내려 새로고침하거나 PWA를 닫았다 다시 여세요.",
    )

    add_divider(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 매뉴얼 끝 —")
    set_korean_font(run, size_pt=10, color=RGBColor(0x99, 0x99, 0x99))

    # 저장
    out_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "계측기-관리-매뉴얼.docx")
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    path = build()
    print(f"매뉴얼 생성 완료: {path}")
